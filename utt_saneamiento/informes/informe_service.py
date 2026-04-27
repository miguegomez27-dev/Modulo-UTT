"""Servicios de generación de informes para UTT Saneamiento."""

from __future__ import annotations

import os
import re
from datetime import datetime

from qgis.PyQt.QtCore import QCoreApplication, QUrl
from qgis.PyQt.QtGui import QDesktopServices

try:
    from docx import Document
    from docx.shared import Cm
except Exception:  # python-docx puede no estar instalado en algunos entornos
    Document = None
    Cm = None


def run_informe_calle(iface, layer, seleccion, ruta_plantilla):
    """Genera informe de calle usando plantilla DOCX y selección actual."""
    if Document is None or Cm is None:
        iface.messageBar().pushWarning(
            "UTT Informes", "Falta dependencia 'python-docx'. No se puede generar el informe."
        )
        return

    directorio_salida = os.path.join(os.path.expanduser("~"), "Documents", "Informes Auto")
    os.makedirs(directorio_salida, exist_ok=True)

    if layer is None:
        iface.messageBar().pushWarning("UTT Informes", "No hay capa para generar informe.")
        return

    if not seleccion:
        iface.messageBar().pushWarning("UTT Informes", "No hay selección de parcelas.")
        return

    if not os.path.exists(ruta_plantilla):
        iface.messageBar().pushCritical("UTT Informes", f"No existe la plantilla: {ruta_plantilla}")
        return

    field_names = set(layer.fields().names())

    def val_feat(feat, campo, default="-"):
        if campo not in field_names:
            return default
        value = feat[campo]
        if value is None:
            return default
        text = str(value).strip()
        if text == "" or text.lower() == "null":
            return default
        return text

    def val_primero_feat(feat, candidatos, default="-"):
        for campo in candidatos:
            out = val_feat(feat, campo, default=None)
            if out not in (None, "", "-"):
                return out
        return default

    def dominio_feat(feat):
        matricula = val_feat(feat, "matricula", "-")
        folio = val_feat(feat, "folio", "-")
        merged = " / ".join([v for v in (matricula, folio) if v != "-"])
        return merged if merged else "-"

    def reemplazar_texto(texto, datos):
        result = texto
        for key, value in datos.items():
            result = re.sub(r"{{\s*" + re.escape(key) + r"\s*}}", str(value), result)
        return result

    def encontrar_tabla_parcelas(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if re.search(r"{{\s*cuenta\s*}}", cell.text):
                        return table
        return None

    def capturar_mapa(ruta_png, margen=1.25):
        canvas = iface.mapCanvas()
        extent = layer.boundingBoxOfSelected()
        if extent.isNull():
            return False

        extent.scale(margen)
        canvas.setExtent(extent)
        canvas.refresh()
        for _ in range(6):
            QCoreApplication.processEvents()

        image = canvas.grab().toImage()
        if image.isNull():
            return False

        return image.save(ruta_png, "PNG")

    def insertar_mapa_en_doc(doc, ruta_png, ancho_cm=16):
        pattern = re.compile(r"{{\s*mapa\s*}}")

        for paragraph in doc.paragraphs:
            if pattern.search(paragraph.text):
                paragraph.text = pattern.sub("", paragraph.text)
                paragraph.add_run().add_picture(ruta_png, width=Cm(ancho_cm))
                return True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if pattern.search(cell.text):
                        cell.text = pattern.sub("", cell.text)
                        cell.paragraphs[0].add_run().add_picture(ruta_png, width=Cm(ancho_cm))
                        return True

        doc.add_paragraph("")
        doc.add_picture(ruta_png, width=Cm(ancho_cm))
        return False

    def reemplazar_mapa_por_texto(doc, texto):
        pattern = re.compile(r"{{\s*mapa\s*}}")
        for paragraph in doc.paragraphs:
            if pattern.search(paragraph.text):
                paragraph.text = pattern.sub(texto, paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if pattern.search(cell.text):
                        cell.text = pattern.sub(texto, cell.text)

    meses = [
        "enero",
        "febrero",
        "marzo",
        "abril",
        "mayo",
        "junio",
        "julio",
        "agosto",
        "septiembre",
        "octubre",
        "noviembre",
        "diciembre",
    ]
    now = datetime.now()
    fecha_texto = f"{now.day} de {meses[now.month - 1]} de {now.year}"

    try:
        cuenta_primera = val_feat(seleccion[0], "cuenta", "sin_cuenta")
        cuenta_safe = str(cuenta_primera).replace("/", "-").replace("\\", "-").strip() or "sin_cuenta"

        doc = Document(ruta_plantilla)
        cuentas_lista = [
            val_feat(feature, "cuenta", "-")
            for feature in seleccion
            if val_feat(feature, "cuenta", "-") != "-"
        ]

        if not cuentas_lista:
            cuentas_txt = "-"
        elif len(cuentas_lista) == 1:
            cuentas_txt = cuentas_lista[0]
        elif len(cuentas_lista) == 2:
            cuentas_txt = f"{cuentas_lista[0]} y {cuentas_lista[1]}"
        else:
            cuentas_txt = ", ".join(cuentas_lista[:-1]) + " y " + cuentas_lista[-1]

        datos_generales = {"fecha_actual": fecha_texto, "cuentas_seleccionadas": cuentas_txt}
        for paragraph in doc.paragraphs:
            replaced = reemplazar_texto(paragraph.text, datos_generales)
            if replaced != paragraph.text:
                paragraph.text = replaced

        tabla = encontrar_tabla_parcelas(doc)
        if tabla is None:
            raise RuntimeError("No se encontró la tabla plantilla con {{cuenta}}.")

        fila_plantilla_idx = None
        for idx, row in enumerate(tabla.rows):
            if any(re.search(r"{{\s*cuenta\s*}}", cell.text) for cell in row.cells):
                fila_plantilla_idx = idx
                break

        if fila_plantilla_idx is None:
            raise RuntimeError("No se encontró la fila plantilla con {{cuenta}}.")

        fila_plantilla = tabla.rows[fila_plantilla_idx]

        for feature in seleccion:
            nueva = tabla.add_row()
            for col in range(len(nueva.cells)):
                nueva.cells[col].text = fila_plantilla.cells[col].text

            datos_fila = {
                "cuenta": val_feat(feature, "cuenta", "-"),
                "par_desig_oficial": val_feat(feature, "par_desig_oficial", "-"),
                "superficie_tierra_rural": val_feat(feature, "superficie_tierra_rural", "-"),
                "dominio": dominio_feat(feature),
                "denominacion_titular": val_primero_feat(
                    feature, ["denomiacion_titular", "denominacion_titular"], "-"
                ),
            }
            for cell in nueva.cells:
                cell.text = reemplazar_texto(cell.text, datos_fila)

        tabla._tbl.remove(tabla.rows[fila_plantilla_idx]._tr)

        try:
            layer.removeSelection()
            layer.select([feat.id() for feat in seleccion])
        except Exception:
            pass

        ruta_png = os.path.join(directorio_salida, f"Mapa_{cuenta_safe}.png")
        if capturar_mapa(ruta_png, margen=1.25):
            insertar_mapa_en_doc(doc, ruta_png, ancho_cm=16)
        else:
            reemplazar_mapa_por_texto(
                doc, "Sin gráfico disponible en el SIT para la/las cuenta/s analizada/s."
            )

        nombre_final = f"Informe_Baja_Calle_{cuenta_safe}.docx"
        ruta_final = os.path.join(directorio_salida, nombre_final)
        doc.save(ruta_final)

        iface.messageBar().pushSuccess("UTT Informes", f"Informe generado: {nombre_final}")
        QDesktopServices.openUrl(QUrl.fromLocalFile(ruta_final))

    except Exception as err:
        iface.messageBar().pushCritical("UTT Informes", f"Fallo al generar informe: {err}")
